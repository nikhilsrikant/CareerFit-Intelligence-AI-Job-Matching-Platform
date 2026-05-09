from __future__ import annotations

import io
import os
import re
import zipfile
from functools import lru_cache
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from typing import Iterable
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "has", "have",
    "in", "into", "is", "it", "of", "on", "or", "that", "the", "this", "to", "with",
    "you", "your", "we", "our", "their", "will", "can", "using", "use", "used", "work",
    "working", "experience", "role", "job", "team", "teams", "company", "skills", "strong",
    "including", "across", "within", "based", "about", "responsibilities", "requirements",
    "required", "preferred", "candidate", "candidates", "ability", "knowledge", "excellent",
    "year", "years", "new", "build", "built", "develop", "developed", "create", "created",
}

# This is intentionally broad. It is not a domain whitelist; it is a starter vocabulary.
# The extractor also discovers unknown skills dynamically from resume sections and websites.
BASE_SKILL_LIBRARY = {
    # Programming / software
    "python", "java", "c++", "cpp", "c", "c#", "javascript", "typescript", "sql", "r", "matlab",
    "go", "golang", "rust", "scala", "bash", "shell", "linux", "git", "github", "api", "apis",
    "rest", "graphql", "fastapi", "flask", "django", "react", "next.js", "node", "node.js",
    "docker", "kubernetes", "terraform", "aws", "gcp", "azure", "postgres", "postgresql",
    "mysql", "mongodb", "redis", "neo4j", "spark", "airflow", "kafka", "etl", "data pipeline",
    # AI / data
    "pandas", "numpy", "scikit-learn", "sklearn", "pytorch", "tensorflow", "keras", "xgboost",
    "lightgbm", "opencv", "computer vision", "object detection", "image segmentation", "deep learning",
    "machine learning", "nlp", "natural language processing", "llm", "large language models", "rag",
    "retrieval", "embeddings", "vector database", "langchain", "llamaindex", "ocr", "analytics",
    "eda", "feature engineering", "a/b testing", "hypothesis testing", "statistics", "shap", "smote",
    "dbscan", "power bi", "tableau", "excel", "looker", "snowflake", "databricks",
    # Electrical / hardware / manufacturing
    "simulink", "labview", "cadence", "altium", "kicad", "ltspice", "pspice", "pcb", "pcb design",
    "circuit design", "analog circuits", "digital logic", "embedded c", "firmware", "microcontrollers",
    "stm32", "arm", "arduino", "raspberry pi", "i2c", "spi", "uart", "can", "fpga", "rtl",
    "verilog", "vhdl", "systemverilog", "asic", "soc", "vlsi", "signal processing", "dsp",
    "control systems", "power electronics", "motor control", "bms", "oscilloscope", "logic analyzer",
    "validation", "verification", "test engineering", "root cause analysis", "six sigma", "lean manufacturing",
    "solidworks", "autocad", "ansys", "catia", "creo", "plc", "scada",
    # Business / product / design / operations / healthcare / finance / marketing / legal
    "product management", "roadmap", "market research", "user research", "figma", "ux", "ui", "wireframes",
    "prototyping", "salesforce", "hubspot", "seo", "sem", "google analytics", "content strategy",
    "financial modeling", "valuation", "forecasting", "accounting", "gaap", "quickbooks", "sap", "erp",
    "clinical research", "patient care", "hipaa", "emr", "epic", "regulatory affairs", "quality assurance",
    "project management", "agile", "scrum", "jira", "confluence", "supply chain", "logistics",
}

ROLE_WORDS = {
    "engineer", "developer", "scientist", "analyst", "designer", "manager", "consultant", "specialist",
    "researcher", "technician", "associate", "coordinator", "architect", "administrator", "operator",
    "strategist", "writer", "editor", "marketer", "accountant", "auditor", "nurse", "therapist",
    "pharmacist", "mechanic", "planner", "recruiter", "lawyer", "paralegal", "teacher", "professor",
    "intern", "trainee", "apprentice", "fellow",
}

TITLE_SECTION_HINTS = re.compile(
    r"(?im)^\s*(professional summary|summary|objective|experience|work experience|technical skills|skills|projects|education|certifications|publications|tools|technologies)\s*$"
)

SKILL_SECTION_HINTS = re.compile(
    r"(?im)^\s*(technical skills|skills|technologies|tools|toolkit|software|languages|certifications|platforms)\s*[:\-]?\s*(.*)$"
)

ROLE_PHRASE_PATTERNS = [
    re.compile(r"(?i)\b(?:seeking|targeting|interested in|looking for|pursuing)\s+(?:a|an)?\s*([^\.\n]{3,90}?)(?:\s+(?:role|position|internship|opportunity))?\b"),
    re.compile(r"(?i)\b([A-Z][A-Za-z/&+\- ]{2,70}\s+(?:Engineer|Developer|Scientist|Analyst|Designer|Manager|Specialist|Consultant|Researcher|Technician|Intern|Architect|Coordinator|Associate))\b"),
]

DOMAIN_HINT_PATTERNS = {
    "AI / Machine Learning": {"machine learning", "deep learning", "model", "pytorch", "tensorflow", "computer vision", "llm", "nlp", "ai", "data science"},
    "Software / Backend": {"software", "backend", "api", "distributed systems", "python", "java", "docker", "kubernetes", "cloud"},
    "Data / Analytics": {"analytics", "sql", "dashboard", "statistics", "tableau", "power bi", "excel", "reporting", "business intelligence"},
    "Electrical / Embedded / Hardware": {"electrical", "electronics", "embedded", "firmware", "pcb", "circuit", "fpga", "rtl", "matlab", "signal", "control", "hardware"},
    "Mechanical / Manufacturing": {"mechanical", "manufacturing", "cad", "solidworks", "ansys", "thermal", "fluid", "quality", "lean", "six sigma"},
    "Product / UX / Design": {"product", "roadmap", "user research", "figma", "ux", "ui", "wireframe", "prototype", "design"},
    "Business / Finance / Operations": {"finance", "financial", "operations", "supply chain", "logistics", "accounting", "forecasting", "strategy"},
    "Healthcare / Life Sciences": {"clinical", "patient", "healthcare", "biology", "biomedical", "pharma", "regulatory", "medical"},
    "Marketing / Growth": {"marketing", "seo", "content", "campaign", "growth", "brand", "social media", "copywriting"},
}


@dataclass
class SourceRecord:
    source_type: str
    name: str
    status: str
    chars: int = 0
    error: str | None = None


@dataclass
class RuntimeProfile:
    combined_text: str = ""
    skills: list[str] = field(default_factory=list)
    keywords: list[str] = field(default_factory=list)
    suggested_roles: list[dict] = field(default_factory=list)
    sources: list[SourceRecord] = field(default_factory=list)
    document_tracks: list[dict] = field(default_factory=list)
    domains: list[dict] = field(default_factory=list)

    def as_match_context(self) -> dict:
        return {
            "combined_text": self.combined_text,
            "skills": self.skills,
            "keywords": self.keywords,
            "suggested_roles": self.suggested_roles,
            "sources": [s.__dict__ for s in self.sources],
            "document_tracks": self.document_tracks,
            "domains": self.domains,
        }


def normalize_url(url: str) -> str:
    url = (url or "").strip()
    if not url:
        return ""
    parsed = urlparse(url)
    if not parsed.scheme:
        url = "https://" + url
    return url


def clean_text(text: str) -> str:
    text = re.sub(r"[\u2022\u25cf\u25aa\u00b7]", " • ", text or " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


@lru_cache(maxsize=128)
def _fetch_public_website_text_cached(url: str, max_chars: int, timeout_seconds: float) -> tuple[str, str | None]:
    headers = {
        "User-Agent": "CareerFitIntelligence/2.2 (+https://example.com; public profile analyzer)",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    try:
        with httpx.Client(timeout=timeout_seconds, follow_redirects=True, headers=headers) as client:
            response = client.get(url)
            response.raise_for_status()
    except Exception as exc:
        return "", str(exc)

    content_type = response.headers.get("content-type", "")
    if "text/html" not in content_type and "application/xhtml" not in content_type:
        return "", f"Unsupported content type: {content_type or 'unknown'}"

    soup = BeautifulSoup(response.text, "html.parser")
    for node in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
        node.decompose()

    parts: list[str] = []
    title = soup.find("title")
    if title and title.get_text(strip=True):
        parts.append(title.get_text(" ", strip=True))

    for attr in ["description", "keywords"]:
        meta = soup.find("meta", attrs={"name": attr})
        if meta and meta.get("content"):
            parts.append(str(meta.get("content")))

    # Keep profile extraction fast: prioritize semantic content tags, then cap total text.
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "a", "strong"]):
        text = tag.get_text(" ", strip=True)
        if text and len(text) > 1:
            parts.append(text)
        if sum(len(x) for x in parts) >= max_chars * 1.2:
            break

    text = clean_text("\n".join(parts))
    return text[:max_chars], None


def fetch_public_website_text(url: str, max_chars: int = 45000) -> tuple[str, str | None]:
    url = normalize_url(url)
    if not url:
        return "", "Empty URL"
    try:
        timeout = float(os.getenv("CAREERFIT_PROFILE_FETCH_TIMEOUT", "10"))
    except ValueError:
        timeout = 10.0
    return _fetch_public_website_text_cached(url, max_chars, timeout)


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    try:
        max_pages = int(os.getenv("CAREERFIT_MAX_PDF_PAGES", "12"))
    except ValueError:
        max_pages = 12
    for page in list(reader.pages)[:max_pages]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    text = clean_text("\n".join(parts))
    try:
        max_chars = int(os.getenv("CAREERFIT_MAX_PROFILE_CHARS", "90000"))
    except ValueError:
        max_chars = 90000
    return text[:max_chars]


def looks_like_pages_file(file_bytes: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = set(zf.namelist())
            return "Index/Document.iwa" in names or any(name.startswith("Index/") and name.endswith(".iwa") for name in names)
    except Exception:
        return False


def extract_docx_text(file_bytes: bytes) -> str:
    if looks_like_pages_file(file_bytes):
        raise ValueError(
            "This file looks like an Apple Pages document renamed as .docx. Export/download it as a real Word .docx or PDF, then upload again."
        )
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = clean_text("\n".join(parts))
    try:
        max_chars = int(os.getenv("CAREERFIT_MAX_PROFILE_CHARS", "90000"))
    except ValueError:
        max_chars = 90000
    return text[:max_chars]


def tokens(text: str) -> list[str]:
    lowered = text.lower()
    raw = re.findall(r"[a-zA-Z][a-zA-Z0-9+.#/-]{1,}", lowered)
    return [t.strip("-/.#") for t in raw if t not in STOPWORDS and len(t.strip("-/.#")) > 2]


def canonical_label(value: str) -> str:
    v = value.strip().strip(",;:|•")
    lower = v.lower()
    mapping = {
        "cpp": "C++", "c++": "C++", "c#": "C#", "sklearn": "scikit-learn", "postgres": "PostgreSQL",
        "postgresql": "PostgreSQL", "next.js": "Next.js", "node.js": "Node.js", "llm": "LLM", "nlp": "NLP",
        "rag": "RAG", "ocr": "OCR", "aws": "AWS", "gcp": "GCP", "api": "API", "apis": "APIs",
        "mlops": "MLOps", "cuda": "CUDA", "gpu": "GPU", "fpga": "FPGA", "rtl": "RTL", "pcb": "PCB",
        "vhdl": "VHDL", "asic": "ASIC", "soc": "SoC", "bms": "BMS", "ui": "UI", "ux": "UX", "seo": "SEO",
        "sem": "SEM", "hipaa": "HIPAA", "gaap": "GAAP", "erp": "ERP", "emr": "EMR", "plc": "PLC", "scada": "SCADA",
    }
    if lower in mapping:
        return mapping[lower]
    if v.isupper() and len(v) <= 6:
        return v
    # Preserve common punctuated tech names.
    if any(ch in v for ch in ["+", "#", "."]):
        return v
    return v.title()


def extract_explicit_skill_sections(text: str) -> set[str]:
    found: set[str] = set()
    lines = re.split(r"[\n\r]+|(?<=\.)\s+", text)
    capture = False
    buffer: list[str] = []
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if SKILL_SECTION_HINTS.match(line):
            capture = True
            buffer.append(SKILL_SECTION_HINTS.sub(r"\2", line))
            continue
        if capture and TITLE_SECTION_HINTS.match(line) and not SKILL_SECTION_HINTS.match(line):
            capture = False
        if capture:
            buffer.append(line)
        if len(buffer) > 16:
            break
    skill_blob = " | ".join(buffer)
    for chunk in re.split(r"[,;|•/]+", skill_blob):
        item = chunk.strip()
        if 2 <= len(item) <= 45 and not item.lower().startswith(("experience", "proficient", "familiar")):
            if not re.match(r"^(and|or|with|including)$", item.lower()):
                found.add(canonical_label(item))
    return found


def extract_capitalized_terms(text: str) -> set[str]:
    found: set[str] = set()
    # Find tool-like terms: C++, MATLAB, AutoCAD, Power BI, Google Analytics, etc.
    patterns = [
        r"\b[A-Z][A-Za-z0-9]*(?:[+#.][A-Za-z0-9]+)+\b",
        r"\b[A-Z]{2,}(?:/[A-Z]{2,})?\b",
        r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2}\b",
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            term = m.group(0).strip()
            if 2 <= len(term) <= 38 and term.lower() not in STOPWORDS:
                # Avoid names/locations by requiring some tech-ish signal or repeated occurrence.
                if term.isupper() or any(ch in term for ch in ["+", "#", "."]) or text.count(term) >= 2:
                    found.add(canonical_label(term))
    return found


def extract_skills(text: str, limit: int = 120) -> list[str]:
    lowered = " " + text.lower() + " "
    found: set[str] = set()
    for skill in BASE_SKILL_LIBRARY:
        pattern = r"(?<![a-z0-9])" + re.escape(skill.lower()) + r"(?![a-z0-9])"
        if re.search(pattern, lowered):
            found.add(canonical_label(skill))
    found |= extract_explicit_skill_sections(text)
    found |= extract_capitalized_terms(text)
    # Filter out very generic resume words.
    bad = {"Education", "Experience", "Projects", "Summary", "Professional", "Technical", "Skills", "Resume"}
    cleaned = [x for x in found if x not in bad and len(x) <= 48]
    return sorted(cleaned, key=lambda x: (x.lower()))[:limit]


def extract_keywords(text: str, limit: int = 120) -> list[str]:
    word_tokens = tokens(text)
    counts = Counter(word_tokens)

    # Dynamic n-grams: no domain restriction.
    for n in (2, 3):
        for i in range(0, max(0, len(word_tokens) - n + 1)):
            phrase_tokens = word_tokens[i:i+n]
            phrase = " ".join(phrase_tokens)
            if any(t in STOPWORDS for t in phrase_tokens):
                continue
            if len(phrase) < 6 or len(phrase) > 55:
                continue
            # Boost phrases that look like capabilities/tools/domains.
            if any(suffix in phrase for suffix in ["engineering", "analysis", "design", "systems", "modeling", "management", "research", "development", "testing", "automation", "operations"]):
                counts[phrase] += 6
            elif counts[phrase_tokens[0]] > 1 and counts[phrase_tokens[-1]] > 1:
                counts[phrase] += 2

    return [word for word, _ in counts.most_common(limit)]


def normalize_role_phrase(phrase: str) -> str:
    phrase = re.sub(r"\s+", " ", phrase).strip(" -:,.|•")
    phrase = re.sub(r"(?i)\b(internship|opportunity|position|role|jobs?)\b", "", phrase).strip(" -:,.|•")
    words = phrase.split()
    if len(words) > 7:
        phrase = " ".join(words[:7])
    return phrase.title()


def extract_explicit_roles(text: str) -> list[dict]:
    roles: dict[str, set[str]] = defaultdict(set)
    for pattern in ROLE_PHRASE_PATTERNS:
        for match in pattern.finditer(text):
            phrase = normalize_role_phrase(match.group(1))
            if not phrase or len(phrase) < 5:
                continue
            lower = phrase.lower()
            if any(bad in lower for bad in ["technical skills", "professional summary", "work experience"]):
                continue
            if any(word in lower.split() for word in ROLE_WORDS) or any(w in lower for w in ["engineering", "analytics", "research", "design", "operations", "finance", "marketing", "clinical"]):
                roles[phrase].add("explicit profile phrase")
    return [{"role": role, "evidence": sorted(evidence), "base": 0.74} for role, evidence in roles.items()]


def infer_domain_roles(text: str, skills: Iterable[str], keywords: Iterable[str]) -> list[dict]:
    haystack = " ".join([text.lower(), " ".join(skills).lower(), " ".join(keywords).lower()])
    results: list[dict] = []
    for domain, signals in DOMAIN_HINT_PATTERNS.items():
        hits = sorted({sig for sig in signals if sig in haystack})
        if not hits:
            continue
        score = min(0.95, 0.42 + 0.055 * len(hits))
        # Convert domain to role suggestions without making the app domain-limited.
        if domain == "AI / Machine Learning":
            role_names = ["Machine Learning Engineer", "AI Engineer", "Data Scientist"]
        elif domain == "Software / Backend":
            role_names = ["Software Engineer", "Backend Engineer", "Platform Engineer"]
        elif domain == "Data / Analytics":
            role_names = ["Data Analyst", "Data Scientist", "Analytics Engineer"]
        elif domain == "Electrical / Embedded / Hardware":
            role_names = ["Electrical Engineer", "Embedded Systems Engineer", "Hardware Engineer", "Firmware Engineer"]
        elif domain == "Mechanical / Manufacturing":
            role_names = ["Mechanical Engineer", "Manufacturing Engineer", "Quality Engineer"]
        elif domain == "Product / UX / Design":
            role_names = ["Product Manager", "UX Designer", "Product Designer"]
        elif domain == "Business / Finance / Operations":
            role_names = ["Business Analyst", "Financial Analyst", "Operations Analyst"]
        elif domain == "Healthcare / Life Sciences":
            role_names = ["Clinical Research Associate", "Biomedical Engineer", "Regulatory Affairs Specialist"]
        elif domain == "Marketing / Growth":
            role_names = ["Marketing Analyst", "Growth Marketer", "Content Strategist"]
        else:
            role_names = [domain]
        for role in role_names:
            results.append({"role": role, "score": round(score, 2), "evidence": hits[:8], "domain": domain})
    return results


def suggest_dynamic_roles(text: str, skills: Iterable[str], keywords: Iterable[str], limit: int = 14) -> list[dict]:
    candidates: dict[str, dict] = {}
    for item in extract_explicit_roles(text):
        role = item["role"]
        candidates[role] = {"role": role, "score": item.get("base", 0.74), "evidence": item.get("evidence", []), "domain": "Profile stated"}
    for item in infer_domain_roles(text, skills, keywords):
        role = item["role"]
        if role not in candidates or item["score"] > candidates[role]["score"]:
            candidates[role] = item
        else:
            candidates[role]["evidence"] = sorted(set(candidates[role].get("evidence", []) + item.get("evidence", [])))[:8]
    # Create role suggestions from keyword phrases ending in role-like words.
    for phrase in keywords:
        p = str(phrase).lower()
        if any(w in p.split() for w in ROLE_WORDS) or any(s in p for s in ["engineering", "analysis", "design", "research", "operations", "management"]):
            role = normalize_role_phrase(str(phrase))
            if 5 <= len(role) <= 60 and role not in candidates:
                candidates[role] = {"role": role, "score": 0.62, "evidence": [str(phrase)], "domain": "Inferred from keywords"}
    results = list(candidates.values())
    results.sort(key=lambda item: (item.get("score", 0), len(item.get("evidence", []))), reverse=True)
    return results[:limit]


def infer_domains(text: str, skills: Iterable[str], keywords: Iterable[str]) -> list[dict]:
    haystack = " ".join([text.lower(), " ".join(skills).lower(), " ".join(keywords).lower()])
    domains = []
    for domain, signals in DOMAIN_HINT_PATTERNS.items():
        hits = sorted({sig for sig in signals if sig in haystack})
        if hits:
            domains.append({"domain": domain, "score": round(min(0.99, len(hits) / max(4, len(signals) * 0.35)), 2), "evidence": hits[:8]})
    domains.sort(key=lambda item: item["score"], reverse=True)
    return domains


def build_runtime_profile(urls: list[str], uploaded_files: list[tuple[str, bytes]]) -> RuntimeProfile:
    sources: list[SourceRecord] = []
    text_parts: list[str] = []
    document_tracks: list[dict] = []

    for url in urls:
        url = normalize_url(url)
        if not url:
            continue
        text, error = fetch_public_website_text(url)
        if error:
            sources.append(SourceRecord("website", url, "error", 0, error))
        elif not text.strip():
            sources.append(SourceRecord("website", url, "error", 0, "No readable public text found"))
        else:
            sources.append(SourceRecord("website", url, "parsed", len(text)))
            text_parts.append(f"Website source: {url}\n{text}")

    for name, data in uploaded_files:
        lowered = name.lower()
        try:
            if lowered.endswith(".pdf"):
                text = extract_pdf_text(data)
            elif lowered.endswith(".docx"):
                text = extract_docx_text(data)
            else:
                sources.append(SourceRecord("document", name, "error", 0, "Only PDF and DOCX are supported"))
                continue
            if not text.strip():
                sources.append(SourceRecord("document", name, "error", 0, "No readable text was extracted. Try exporting as PDF or true Word .docx."))
                continue
            sources.append(SourceRecord("document", name, "parsed", len(text)))
            text_parts.append(f"Document source: {name}\n{text}")
            doc_skills = extract_skills(text)
            doc_keywords = extract_keywords(text, 60)
            document_tracks.append({
                "name": name,
                "chars": len(text),
                "skills": doc_skills[:45],
                "keywords": doc_keywords[:55],
                "suggested_roles": suggest_dynamic_roles(text, doc_skills, doc_keywords, 8),
            })
        except Exception as exc:
            sources.append(SourceRecord("document", name, "error", 0, str(exc)))

    combined = clean_text("\n\n".join(text_parts))
    skills = extract_skills(combined)
    keywords = extract_keywords(combined, 140)
    roles = suggest_dynamic_roles(combined, skills, keywords, 16)
    domains = infer_domains(combined, skills, keywords)
    return RuntimeProfile(combined, skills, keywords, roles, sources, document_tracks, domains)
